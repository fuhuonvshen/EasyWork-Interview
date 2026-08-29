"""Markdown → DOCX / PDF / PNG rendering.

Extracted from main.py to keep the chat logic separate from document rendering.
"""

import base64
import logging
import os
import re
from io import BytesIO
from lxml import etree

from .data.models import ExportResponse

logger = logging.getLogger("agent.export")

# ── CJK font discovery ──────────────────────────────────────

_FONT_PATHS = [
    os.path.join(os.environ.get("WINDIR", "C:/Windows"), "Fonts", "simhei.ttf"),
    os.path.join(os.environ.get("WINDIR", "C:/Windows"), "Fonts", "msyh.ttc"),
    os.path.join(os.environ.get("WINDIR", "C:/Windows"), "Fonts", "simsun.ttc"),
    "simhei.ttf",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
]


def _find_cjk_font() -> str:
    for p in _FONT_PATHS:
        if os.path.exists(p):
            return p
    return "simhei.ttf"


def _parse_inline(text: str) -> list:
    """Parse **bold**, *italic*, `code` into list of (text, bold, italic, code)."""
    segments = []
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False, False, False))
        if m.group(2) is not None:
            segments.append((m.group(2), True, False, False))
        elif m.group(3) is not None:
            segments.append((m.group(3), False, True, False))
        elif m.group(4) is not None:
            segments.append((m.group(4), False, False, True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False, False, False))
    return segments if segments else [(text, False, False, False)]


def _to_xml(text: str) -> str:
    """Convert markdown inline to reportlab XML with escaped entities."""
    result = []
    for seg_text, bold, italic, code in _parse_inline(text):
        escaped = seg_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if code:
            result.append(f'<font face="Courier" color="#333333">{escaped}</font>')
        elif bold and italic:
            result.append(f"<b><i>{escaped}</i></b>")
        elif bold:
            result.append(f"<b>{escaped}</b>")
        elif italic:
            result.append(f"<i>{escaped}</i>")
        else:
            result.append(escaped)
    return "".join(result)


async def render_export(content: str, fmt: str) -> ExportResponse:
    """Convert markdown content to the requested format.

    Returns an ExportResponse with base64-encoded data and a filename.
    """
    logger.info("[export] format=%s content_len=%d", fmt, len(content))

    CJK_FONT_PATH = _find_cjk_font()
    CJK_FONT_NAME = "SimHei"
    FALLBACK_FONT = "Helvetica"
    logger.info("[export] using CJK font: %s", CJK_FONT_PATH)

    # ── DOCX ──
    if fmt == "docx":
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        rpr = style.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rpr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), "SimHei")

        def _set_run_font(run, font_name="SimHei"):
            run.font.name = font_name
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = etree.SubElement(rPr, qn("w:rFonts"))
            rFonts.set(qn("w:eastAsia"), font_name)

        for line in content.split("\n"):
            stripped = line.rstrip()
            if stripped.startswith("### "):
                h = doc.add_heading(stripped[4:], level=3)
                for run in h.runs:
                    _set_run_font(run)
            elif stripped.startswith("## "):
                h = doc.add_heading(stripped[3:], level=2)
                for run in h.runs:
                    _set_run_font(run)
            elif stripped.startswith("# "):
                h = doc.add_heading(stripped[2:], level=1)
                for run in h.runs:
                    _set_run_font(run)
            elif re.match(r'^[\s]*[-*]\s+', stripped):
                item_text = re.sub(r'^[\s]*[-*]\s+', '', stripped)
                p = doc.add_paragraph(style="List Bullet")
                p.clear()
                for seg_text, bold, italic, code in _parse_inline(item_text):
                    run = p.add_run(seg_text)
                    run.bold = bold
                    run.italic = italic
                    _set_run_font(run, "Courier New" if code else "SimHei")
            elif re.match(r'^[\s]*\d+[.)]\s+', stripped):
                item_text = re.sub(r'^[\s]*\d+[.)]\s+', '', stripped)
                p = doc.add_paragraph(style="List Number")
                p.clear()
                for seg_text, bold, italic, code in _parse_inline(item_text):
                    run = p.add_run(seg_text)
                    run.bold = bold
                    run.italic = italic
                    _set_run_font(run, "Courier New" if code else "SimHei")
            elif not stripped:
                doc.add_paragraph()
            else:
                p = doc.add_paragraph()
                for seg_text, bold, italic, code in _parse_inline(stripped):
                    run = p.add_run(seg_text)
                    run.bold = bold
                    run.italic = italic
                    _set_run_font(run, "Courier New" if code else "SimHei")

        buf = BytesIO()
        doc.save(buf)
        data = base64.b64encode(buf.getvalue()).decode()
        return ExportResponse(data=data, filename="报告.docx")

    # ── PDF ──
    if fmt == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT

        try:
            if not CJK_FONT_PATH:
                raise FileNotFoundError("CJK font not found on system")
            pdfmetrics.registerFont(TTFont(CJK_FONT_NAME, CJK_FONT_PATH))
            logger.info("[export] PDF registered font: %s from %s", CJK_FONT_NAME, CJK_FONT_PATH)
        except Exception as e:
            logger.warning("[export] PDF font registration failed: %s — falling back to %s", e, FALLBACK_FONT)
            CJK_FONT_NAME = FALLBACK_FONT

        buf = BytesIO()
        pdf = SimpleDocTemplate(buf, pagesize=A4,
                                topMargin=20 * mm, bottomMargin=20 * mm,
                                leftMargin=20 * mm, rightMargin=20 * mm)
        base_style = ParagraphStyle("CJKBase", fontName=CJK_FONT_NAME, fontSize=11,
                                    leading=18, spaceAfter=6, alignment=TA_LEFT)
        h1_style = ParagraphStyle("CJKH1", parent=base_style, fontSize=20, leading=28,
                                  spaceBefore=16, spaceAfter=10)
        h2_style = ParagraphStyle("CJKH2", parent=base_style, fontSize=16, leading=22,
                                  spaceBefore=12, spaceAfter=8)
        h3_style = ParagraphStyle("CJKH3", parent=base_style, fontSize=13, leading=18,
                                  spaceBefore=8, spaceAfter=6)
        bullet_style = ParagraphStyle("CJKBullet", parent=base_style, leftIndent=24,
                                      bulletIndent=12, spaceBefore=1, spaceAfter=1)
        story = []
        for line in content.split("\n"):
            stripped = line.rstrip()
            if stripped.startswith("### "):
                story.append(Paragraph(_to_xml(stripped[4:]), h3_style))
            elif stripped.startswith("## "):
                story.append(Paragraph(_to_xml(stripped[3:]), h2_style))
            elif stripped.startswith("# "):
                story.append(Paragraph(_to_xml(stripped[2:]), h1_style))
            elif re.match(r'^[\s]*[-*]\s+', stripped):
                item_text = re.sub(r'^[\s]*[-*]\s+', '', stripped)
                story.append(Paragraph(f"· {_to_xml(item_text)}", bullet_style))
            elif re.match(r'^[\s]*\d+[.)]\s+', stripped):
                item_text = re.sub(r'^[\s]*\d+[.)]\s+', '', stripped)
                story.append(Paragraph(_to_xml(item_text), bullet_style))
            elif stripped:
                story.append(Paragraph(_to_xml(stripped), base_style))
            else:
                story.append(Spacer(1, 6))
        pdf.build(story)
        data = base64.b64encode(buf.getvalue()).decode()
        return ExportResponse(data=data, filename="报告.pdf")

    # ── PNG ──
    if fmt == "png":
        from PIL import Image, ImageDraw, ImageFont

        def _load_font(size: int) -> ImageFont.FreeTypeFont:
            try:
                return ImageFont.truetype(CJK_FONT_PATH, size)
            except Exception:
                return ImageFont.load_default()

        font_normal = _load_font(16)
        font_h1 = _load_font(24)
        font_h2 = _load_font(20)
        font_h3 = _load_font(17)
        font_small = _load_font(14)
        margin = 40
        img_width = 900
        max_text_width = img_width - margin * 2

        def _text_width(text: str, font) -> int:
            try:
                bbox = font.getbbox(text)
                return bbox[2] - bbox[0]
            except Exception:
                return len(text) * (font.size // 2)

        draw_lines: list = []
        for line in content.split("\n"):
            stripped = line.rstrip()
            if not stripped:
                draw_lines.append(("", font_normal, 4, "black"))
                continue
            line_font = font_normal
            color = "black"
            extra_space = 0
            if stripped.startswith("### "):
                text = stripped[4:]
                line_font = font_h3
                extra_space = 6
            elif stripped.startswith("## "):
                text = stripped[3:]
                line_font = font_h2
                extra_space = 8
            elif stripped.startswith("# "):
                text = stripped[2:]
                line_font = font_h1
                extra_space = 10
            elif re.match(r'^[\s]*[-*]\s+', stripped):
                text = "  · " + re.sub(r'^[\s]*[-*]\s+', '', stripped)
            elif re.match(r'^[\s]*\d+[.)]\s+', stripped):
                text = "  " + re.sub(r'^[\s]*', '', stripped)
            elif re.match(r'^[-*_]{3,}$', stripped):
                draw_lines.append(("─" * 60, font_small, 8, "#cccccc"))
                continue
            else:
                text = stripped
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            current = ""
            for ch in text:
                test = current + ch
                if _text_width(test, line_font) > max_text_width and current:
                    draw_lines.append((current, line_font, extra_space, color))
                    current = ch
                    extra_space = 0
                else:
                    current = test
            if current:
                draw_lines.append((current, line_font, extra_space, color))

        MAX_PNG_HEIGHT = 5000
        total_h = margin
        for _, f, ext, _ in draw_lines:
            total_h += (f.size + 8) + ext
        total_h += margin
        if total_h > MAX_PNG_HEIGHT:
            logger.warning("[export] PNG content too tall (%dpx), truncating to %dpx", total_h, MAX_PNG_HEIGHT)
        canvas_h = max(200, min(total_h, MAX_PNG_HEIGHT))
        img = Image.new("RGB", (img_width, canvas_h), "white")
        draw = ImageDraw.Draw(img)
        y = margin
        for line_text, f, ext, color in draw_lines:
            if y > canvas_h - 40:
                draw.text((margin, canvas_h - 30), "... (内容过长，已截断)", fill="#999999", font=font_small)
                break
            y += ext
            if line_text:
                draw.text((margin, y), line_text, fill=color, font=f)
            y += f.size + 8
        buf = BytesIO()
        img.save(buf, format="PNG")
        data = base64.b64encode(buf.getvalue()).decode()
        return ExportResponse(data=data, filename="报告.png")

    # ── Markdown (fallback) ──
    return ExportResponse(data=base64.b64encode(content.encode()).decode(), filename="报告.md")
